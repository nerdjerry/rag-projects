"""
document_parser.py — Parse legal documents (PDF and DOCX) while preserving structure.

WHY STRUCTURE MATTERS FOR LEGAL DOCUMENTS:
    Legal contracts are highly structured. Section numbers (e.g., "3.2 Termination")
    carry legal significance — courts reference them, lawyers cite them, and automated
    analysis needs them to attribute clauses correctly. Losing structure during parsing
    means losing the ability to say "Based on Section 4.1: ..." in our answers.

    Unlike casual text, legal documents use numbering, headings, and defined terms
    in very specific ways. Preserving these lets us:
    1. Cite exact sections in Q&A answers (critical for legal credibility).
    2. Group related clauses for conflict detection.
    3. Identify missing sections (e.g., no "Limitation of Liability" heading).
"""

import os
import re
from typing import List

from langchain.schema import Document


# ---------------------------------------------------------------------------
# Patterns that typically indicate a major section heading in a legal document.
# Examples matched: "1. DEFINITIONS", "3.2 Termination Rights", "ARTICLE IV",
#                   "SECTION 5 — INDEMNIFICATION"
# ---------------------------------------------------------------------------
SECTION_HEADING_PATTERN = re.compile(
    r"^(?:"
    r"(?:ARTICLE|SECTION|EXHIBIT)\s+[IVXLCDM\d]+"          # ARTICLE IV, SECTION 5
    r"|(?:\d+\.)+\s*[A-Z]"                                  # 3.2 Termination
    r"|[A-Z][A-Z\s]{4,}"                                    # ALL-CAPS HEADING
    r")",
    re.MULTILINE,
)

# Common legal section labels used for tagging chunks with their topic.
LEGAL_SECTIONS = {
    "parties":      ["parties", "between", "by and between", "recitals"],
    "definitions":  ["definitions", "defined terms", "interpretation"],
    "term":         ["term", "duration", "effective date", "commencement"],
    "obligations":  ["obligations", "responsibilities", "duties", "covenants"],
    "payment":      ["payment", "compensation", "fees", "invoic"],
    "termination":  ["termination", "expiration", "cancellation"],
    "liability":    ["liability", "limitation of liability", "damages"],
    "indemnification": ["indemnif"],
    "confidentiality": ["confidential", "non-disclosure", "nda"],
    "ip":           ["intellectual property", "ip ownership", "work product"],
    "governing_law": ["governing law", "jurisdiction", "dispute resolution"],
}


def _classify_section(text: str) -> str:
    """Return a section label (e.g., 'termination') based on keyword matching.

    This is a simple heuristic — production systems would use a trained classifier,
    but keyword matching works surprisingly well for standard contract language.
    """
    text_lower = text.lower()
    for label, keywords in LEGAL_SECTIONS.items():
        for kw in keywords:
            if kw in text_lower:
                return label
    return "general"


def _extract_section_number(text: str) -> str:
    """Try to pull the section number from the beginning of a text block.

    Returns the matched number string (e.g., '3.2') or an empty string.
    """
    match = re.match(r"^((?:\d+\.)+\d*)\s", text)
    if match:
        return match.group(1).rstrip(".")
    return ""


# ---------------------------------------------------------------------------
# PDF parsing
# ---------------------------------------------------------------------------

def _parse_pdf(file_path: str) -> List[Document]:
    """Parse a PDF file into a list of Document objects.

    Each page becomes one Document. We store the page number and any detected
    section heading in the metadata so downstream components can cite sources.
    """
    # pypdf is a pure-Python PDF reader — no system dependencies required.
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    documents: List[Document] = []

    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if not text.strip():
            continue

        # Try to detect the first section heading on this page.
        heading_match = SECTION_HEADING_PATTERN.search(text)
        section_heading = heading_match.group(0).strip() if heading_match else ""

        section_number = _extract_section_number(text)
        section_label = _classify_section(text)

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "source": os.path.basename(file_path),
                    "page": page_num,
                    "section_heading": section_heading,
                    "section_number": section_number,
                    "section_label": section_label,
                    "file_type": "pdf",
                },
            )
        )

    return documents


# ---------------------------------------------------------------------------
# DOCX parsing
# ---------------------------------------------------------------------------

def _parse_docx(file_path: str) -> List[Document]:
    """Parse a DOCX file into a list of Document objects.

    DOCX files expose paragraph styles (e.g., 'Heading 1') which are even more
    reliable than regex for detecting section boundaries.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    documents: List[Document] = []
    current_heading = ""
    current_text_lines: List[str] = []
    chunk_index = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # python-docx exposes the Word style name — headings are "Heading 1", etc.
        style_name = (para.style.name or "").lower()
        is_heading = "heading" in style_name or text.isupper()

        if is_heading and current_text_lines:
            # Flush the accumulated text as one Document before starting a new section.
            combined = "\n".join(current_text_lines)
            section_number = _extract_section_number(combined)
            section_label = _classify_section(combined)

            documents.append(
                Document(
                    page_content=combined,
                    metadata={
                        "source": os.path.basename(file_path),
                        "chunk_index": chunk_index,
                        "section_heading": current_heading,
                        "section_number": section_number,
                        "section_label": section_label,
                        "file_type": "docx",
                    },
                )
            )
            chunk_index += 1
            current_text_lines = []

        if is_heading:
            current_heading = text
        current_text_lines.append(text)

    # Don't forget the last section.
    if current_text_lines:
        combined = "\n".join(current_text_lines)
        section_number = _extract_section_number(combined)
        section_label = _classify_section(combined)

        documents.append(
            Document(
                page_content=combined,
                metadata={
                    "source": os.path.basename(file_path),
                    "chunk_index": chunk_index,
                    "section_heading": current_heading,
                    "section_number": section_number,
                    "section_label": section_label,
                    "file_type": "docx",
                },
            )
        )

    return documents


# ---------------------------------------------------------------------------
# Plain-text fallback
# ---------------------------------------------------------------------------

def _parse_txt(file_path: str) -> List[Document]:
    """Fallback parser for .txt files — splits on blank-line boundaries."""
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()

    # Split on double newlines to approximate paragraph boundaries.
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]

    documents: List[Document] = []
    for idx, para in enumerate(paragraphs):
        section_label = _classify_section(para)
        documents.append(
            Document(
                page_content=para,
                metadata={
                    "source": os.path.basename(file_path),
                    "chunk_index": idx,
                    "section_label": section_label,
                    "file_type": "txt",
                },
            )
        )
    return documents


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_legal_document(file_path: str) -> List[Document]:
    """Parse a legal document and return structured Document objects.

    Supported formats: .pdf, .docx, .txt

    Each Document's metadata includes:
        - source:          Original filename
        - section_heading: Detected heading text (if any)
        - section_number:  Detected numbering like '3.2' (if any)
        - section_label:   Classified topic (e.g., 'termination', 'liability')
        - file_type:       The format that was parsed

    Args:
        file_path: Path to the legal document file.

    Returns:
        A list of langchain Document objects with enriched metadata.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError:        If the file format is not supported.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Document not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _parse_pdf(file_path)
    elif ext == ".docx":
        return _parse_docx(file_path)
    elif ext == ".txt":
        return _parse_txt(file_path)
    else:
        raise ValueError(
            f"Unsupported file format: '{ext}'. Supported formats: .pdf, .docx, .txt"
        )


# ---------------------------------------------------------------------------
# Quick smoke test — run this file directly to test with a sample document.
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python document_parser.py <path-to-document>")
        sys.exit(1)

    docs = parse_legal_document(sys.argv[1])
    print(f"\nParsed {len(docs)} document chunks from '{sys.argv[1]}':\n")
    for i, doc in enumerate(docs):
        label = doc.metadata.get("section_label", "?")
        heading = doc.metadata.get("section_heading", "")
        preview = doc.page_content[:120].replace("\n", " ")
        print(f"  [{i}] ({label}) {heading}")
        print(f"       {preview}...")
        print()
